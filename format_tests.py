#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Генерация переформатированных docx/txt со сквозной нумерацией.

Формат на вопрос:
    N. <текст вопроса>
    <буква>. <текст правильного ответа>
"""
import os
import re
import docx
from docx.shared import Pt
from parse_tests import parse

OUT_DIR = 'formatted'

FILES = [
    '59.docx',
    '60-123.docx',
    '123-182.docx',
    '182-258.docx',
    '182-258 (2).docx',
    '259-323.docx',
    '259-332.docx',
    '333-400.docx',
    '401-461.docx',
]


def clean_q(text):
    # убираем ведущие дефисы/точки/пробелы и схлопываем пробелы
    text = re.sub(r'^[\s\-–—.•]+', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    # убираем просочившийся стартовый номер диапазона (например "259 ", "333 "),
    # но только большой (>=50), чтобы не задеть осмысленные числа вроде "14 ёш"
    m = re.match(r'^(\d{2,4})\s+(\S)', text)
    if m and int(m.group(1)) >= 50:
        text = text[m.start(2):]
    return text


def clean_a(text):
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def build_docx(items, path):
    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    for i, q in enumerate(items, 1):
        pq = doc.add_paragraph()
        run = pq.add_run(f"{i}. {clean_q(q['q'])}")
        run.bold = True
        letter = q['ans_letter'] or '?'
        doc.add_paragraph(f"{letter}. {clean_a(q['ans_text'])}")
        doc.add_paragraph('')  # пустая строка между вопросами
    doc.save(path)


def build_txt(items, path):
    lines = []
    for i, q in enumerate(items, 1):
        letter = q['ans_letter'] or '?'
        lines.append(f"{i}. {clean_q(q['q'])}")
        lines.append(f"{letter}. {clean_a(q['ans_text'])}")
        lines.append('')
    with open(path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    total = 0
    summary = []
    for f in FILES:
        if not os.path.exists(f):
            print(f'ПРОПУЩЕН (нет файла): {f}')
            continue
        items = parse(f)
        base = os.path.splitext(f)[0]
        build_docx(items, os.path.join(OUT_DIR, f))
        build_txt(items, os.path.join(OUT_DIR, base + '.txt'))
        total += len(items)
        summary.append((f, len(items)))
        print(f'{f:22s} -> {len(items)} вопросов')
    print('-' * 40)
    print(f'ИТОГО вопросов: {total}')
    return summary


if __name__ == '__main__':
    main()
